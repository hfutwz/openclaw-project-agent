#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 paper_final.md 转换为符合 GB/T 7713.2-2022 学术论文格式的 Word 文档

格式规范依据：
  - GB/T 7713.2-2022《学术论文编写规则》
  - GB/T 7714-2015《信息与文献 参考文献著录规则》
  - 《计算机学报》投稿规范

排版规格：
  - 页面：A4，上下页边距 2.54cm，左右页边距 3.17cm
  - 正文字体：中文宋体 / 英文 Times New Roman，小四号（12pt）
  - 行距：固定值 22pt（约 1.5 倍）
  - 章节标题：一级 黑体三号居中；二级 黑体四号；三级 楷体小四号
  - 摘要关键词：摘要标题黑体小四；正文宋体小四
  - 参考文献：宋体五号（10.5pt）
  - 代码块：Courier New 9pt，带浅灰背景底纹
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ───────────────────────────────────────────
# 0. 读取源文件
# ───────────────────────────────────────────
SRC = Path(__file__).parent / "paper_final.md"
OUT = Path(__file__).parent / "paper_final.docx"
text = SRC.read_text(encoding="utf-8")

# ───────────────────────────────────────────
# 1. 文档对象 & 页面设置
# ───────────────────────────────────────────
doc = Document()

# 页面尺寸 A4
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(3.17)
section.right_margin  = Cm(3.17)

# ───────────────────────────────────────────
# 2. 工具函数
# ───────────────────────────────────────────
def set_run_font(run, zh_font="宋体", en_font="Times New Roman", size_pt=12, bold=False, italic=False):
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    # 中文字体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), zh_font)

def set_para_format(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=0, space_after=0,
                    line_rule=WD_LINE_SPACING.EXACTLY, line_val=22,
                    first_line_indent=0):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = line_rule
    pf.line_spacing = Pt(line_val)
    if first_line_indent:
        pf.first_line_indent = Pt(first_line_indent)

def add_shading(cell, fill="F2F2F2"):
    """给表格单元格添加底纹"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def add_para_shading(para, fill="F2F2F2"):
    """给段落添加底纹（用于代码块）"""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def add_border(para):
    """给段落加左侧竖线边框（代码块风格）"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "AAAAAA")
    pBdr.append(left)
    pPr.append(pBdr)

def render_inline(para, text_str, base_zh="宋体", base_en="Times New Roman",
                  base_pt=12, base_bold=False):
    """
    解析行内 **bold** 和 `code` 标记，分段添加 run
    """
    # 先处理 **text**
    pattern = re.compile(r'\*\*(.+?)\*\*|`([^`]+)`|(.[^*`]*)')
    pos = 0
    for m in re.finditer(r'\*\*(.+?)\*\*|`([^`]+)`', text_str):
        # 普通文字（m 之前）
        normal = text_str[pos:m.start()]
        if normal:
            run = para.add_run(normal)
            set_run_font(run, base_zh, base_en, base_pt, base_bold)
        if m.group(1):   # **bold**
            run = para.add_run(m.group(1))
            set_run_font(run, base_zh, base_en, base_pt, True)
        elif m.group(2): # `code`
            run = para.add_run(m.group(2))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = m.end()
    # 尾部剩余
    tail = text_str[pos:]
    if tail:
        run = para.add_run(tail)
        set_run_font(run, base_zh, base_en, base_pt, base_bold)

# ───────────────────────────────────────────
# 3. 标题页（题名 + 摘要 + 关键词）
# ───────────────────────────────────────────
def add_title_block(doc, title_lines, abstract, keywords):
    # 主标题
    main_title = title_lines[0].lstrip("#").strip()
    sub_title = title_lines[1].lstrip("#—— ").strip() if len(title_lines) > 1 else ""

    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=12, space_after=6, line_val=28)
    run = p.add_run(main_title)
    set_run_font(run, zh_font="黑体", en_font="Times New Roman",
                 size_pt=16, bold=True)

    if sub_title:
        p2 = doc.add_paragraph()
        set_para_format(p2, align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=12, line_val=22)
        run2 = p2.add_run("——" + sub_title)
        set_run_font(run2, zh_font="黑体", en_font="Times New Roman",
                     size_pt=14, bold=True)

    # 摘要标题
    p_abs_title = doc.add_paragraph()
    set_para_format(p_abs_title, space_before=6, space_after=2, line_val=22)
    r = p_abs_title.add_run("摘  要：")
    set_run_font(r, "黑体", "Times New Roman", 12, bold=True)
    r2 = p_abs_title.add_run(abstract)
    set_run_font(r2, "宋体", "Times New Roman", 12)

    # 关键词
    p_kw = doc.add_paragraph()
    set_para_format(p_kw, space_before=2, space_after=6, line_val=22)
    rk = p_kw.add_run("关键词：")
    set_run_font(rk, "黑体", "Times New Roman", 12, bold=True)
    rk2 = p_kw.add_run(keywords)
    set_run_font(rk2, "宋体", "Times New Roman", 12)

    # 分隔线
    doc.add_paragraph("─" * 40)

# ───────────────────────────────────────────
# 4. 解析 Markdown 并渲染
# ───────────────────────────────────────────
lines = text.split("\n")

# 提取标题行（# 和 ##开头的前几行作为标题页）
title_lines = []
abstract_text = ""
keywords_text = ""
body_start = 0

# 解析前置信息
i = 0
while i < len(lines):
    line = lines[i].strip()
    if line.startswith("# "):
        title_lines.append(line)
        i += 1
    elif line.startswith("**摘要**"):
        abstract_text = re.sub(r'\*\*摘要\*\*[：:]', '', line).strip()
        i += 1
    elif line.startswith("**关键词**"):
        keywords_text = re.sub(r'\*\*关键词\*\*[：:]', '', line).strip()
        i += 1
        body_start = i
        break
    else:
        if title_lines:
            body_start = i
            break
        i += 1

# 如果没有找到关键词行则重新查找
if not abstract_text:
    for idx, l in enumerate(lines):
        if l.startswith("**摘要**"):
            abstract_text = re.sub(r'\*\*摘要\*\*[：:]', '', l).strip()
        if l.startswith("**关键词**"):
            keywords_text = re.sub(r'\*\*关键词\*\*[：:]', '', l).strip()
            body_start = idx + 1
            break

add_title_block(doc, title_lines, abstract_text, keywords_text)

# ───────────────────────────────────────────
# 5. 正文渲染状态机
# ───────────────────────────────────────────
in_code_block = False
code_lines = []
in_table = False
table_rows = []
table_obj = None

def flush_code(doc, code_lines):
    """将代码块写入文档"""
    if not code_lines:
        return
    for cl in code_lines:
        p = doc.add_paragraph()
        set_para_format(p, space_before=0, space_after=0,
                        line_rule=WD_LINE_SPACING.SINGLE, line_val=14)
        add_para_shading(p, "F5F5F5")
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        run = p.add_run(cl)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def flush_table(doc, rows):
    """将 Markdown 表格写入文档"""
    if not rows:
        return
    # 过滤分隔行
    data_rows = [r for r in rows if not re.match(r'^[\|\s\-:]+$', r)]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    if not parsed:
        return

    col_count = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=col_count)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(parsed):
        for ci, cell_text in enumerate(row):
            if ci >= col_count:
                break
            cell = tbl.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.clear()
            set_para_format(p, space_before=1, space_after=1,
                            line_rule=WD_LINE_SPACING.SINGLE, line_val=16)
            # 表头加粗
            is_header = (ri == 0)
            # 去除 markdown 粗体标记
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            run = p.add_run(clean)
            set_run_font(run, "宋体", "Times New Roman", 10.5, bold=is_header)
            if is_header:
                add_shading(cell, "D9E1F2")

    # 表格后空一行
    doc.add_paragraph()

i = body_start
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # ── 代码块 ──
    if stripped.startswith("```"):
        if not in_code_block:
            in_code_block = True
            code_lines = []
        else:
            in_code_block = False
            flush_code(doc, code_lines)
            # 代码块后加空行
            p = doc.add_paragraph()
            set_para_format(p, space_before=0, space_after=4, line_val=14)
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # ── 表格 ──
    if stripped.startswith("|"):
        in_table = True
        table_rows.append(stripped)
        i += 1
        continue
    else:
        if in_table:
            in_table = False
            flush_table(doc, table_rows)
            table_rows = []

    # ── 跳过水平分隔线 ──
    if re.match(r'^[-─]{3,}$', stripped):
        i += 1
        continue

    # ── 空行 ──
    if not stripped:
        i += 1
        continue

    # ── 一级标题 (## 一、...) ──
    if re.match(r'^## [一二三四五六七八九十]', stripped) or stripped == "## 参考文献" or stripped == "## 附录":
        p = doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=18, space_after=6, line_val=28)
        title_text = re.sub(r'^##\s*', '', stripped)
        run = p.add_run(title_text)
        set_run_font(run, "黑体", "Times New Roman", 14, bold=True)
        i += 1
        continue

    # ── 二级标题 (### 2.x / 附录一/附录二) ──
    if stripped.startswith("### ") or re.match(r'^### 附录[一二]', stripped):
        p = doc.add_paragraph()
        set_para_format(p, space_before=10, space_after=4, line_val=22)
        title_text = re.sub(r'^###\s*', '', stripped)
        run = p.add_run(title_text)
        set_run_font(run, "黑体", "Times New Roman", 12, bold=True)
        i += 1
        continue

    # ── 四级标题 (#### 内容) ──
    if stripped.startswith("#### "):
        p = doc.add_paragraph()
        set_para_format(p, space_before=6, space_after=2, line_val=22)
        title_text = re.sub(r'^####\s*', '', stripped)
        run = p.add_run(title_text)
        set_run_font(run, "楷体", "Times New Roman", 12, bold=False)
        i += 1
        continue

    # ── 参考文献条目 [1] ... ──
    if re.match(r'^\[\d+\]', stripped):
        p = doc.add_paragraph()
        set_para_format(p, space_before=1, space_after=1,
                        line_rule=WD_LINE_SPACING.SINGLE, line_val=18)
        pf = p.paragraph_format
        pf.left_indent = Pt(21)
        pf.first_line_indent = Pt(-21)
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
        run = p.add_run(clean)
        set_run_font(run, "宋体", "Times New Roman", 10.5)
        i += 1
        continue

    # ── 无序列表 - ──
    if re.match(r'^[-*]\s+', stripped):
        p = doc.add_paragraph()
        set_para_format(p, space_before=1, space_after=1, line_val=20)
        pf = p.paragraph_format
        pf.left_indent = Pt(21)
        pf.first_line_indent = Pt(-12)
        item_text = re.sub(r'^[-*]\s+', '• ', stripped)
        render_inline(p, item_text)
        i += 1
        continue

    # ── 有序列表 1. ──
    if re.match(r'^\d+\.\s+', stripped):
        p = doc.add_paragraph()
        set_para_format(p, space_before=1, space_after=1, line_val=20)
        pf = p.paragraph_format
        pf.left_indent = Pt(21)
        pf.first_line_indent = Pt(-14)
        render_inline(p, stripped)
        i += 1
        continue

    # ── 正文段落 ──
    p = doc.add_paragraph()
    set_para_format(p, space_before=0, space_after=4, line_val=22,
                    first_line_indent=24)
    render_inline(p, stripped)
    i += 1

# 最后清理残留表格
if table_rows:
    flush_table(doc, table_rows)

# ───────────────────────────────────────────
# 6. 保存
# ───────────────────────────────────────────
doc.save(str(OUT))
print(f"✅ 已生成：{OUT}")
print(f"   共 {len(doc.paragraphs)} 个段落")
